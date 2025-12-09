import streamlit as st
import sys
import os
import json
import pandas as pd
from datetime import datetime

# Add project root to path
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
sys.path.insert(0, project_root)

from app.api.frontend_api import (
    list_rubric_sets,
    extract_and_save_rubric_from_pdf,
    list_submissions_for_rubric
)

from app.core.config import get_ist_timezone
from frontend.pages.utils.session_manager import check_authentication

st.set_page_config(page_title="Report Evaluation", page_icon="📝", layout="wide")

# Check authentication
check_authentication('teacher')

# MSRIT Color Scheme
st.markdown("""
<style>
    /* MSRIT Color Theme */
    .stApp {
        background-color: #f5f5f5;
    }
    
    /* Primary buttons - Red */
    .stButton > button[kind="primary"] {
        background-color: #c41e3a !important;
        color: white !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background-color: #a01830 !important;
    }
    
    /* Regular buttons - Navy */
    .stButton > button {
        background-color: #2d3e50 !important;
        color: white !important;
    }
    
    .stButton > button:hover {
        background-color: #1f2d3d !important;
    }
    
    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: #2d3e50 !important;
    }
    
    section[data-testid="stSidebar"] * {
        color: white !important;
    }
    
    /* Headers */
    h1, h2, h3 {
        color: #2d3e50 !important;
    }
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        background-color: #2d3e50;
    }
    
    .stTabs [data-baseweb="tab"] {
        color: white !important;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: #c41e3a !important;
    }
    
    /* Form elements */
    .stTextInput > label, .stDateInput > label, .stTimeInput > label, 
    .stNumberInput > label, .stFileUploader > label {
        color: #2d3e50 !important;
        font-weight: 600 !important;
    }
    
    /* Dataframes */
    .stDataFrame {
        border: 2px solid #2d3e50 !important;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background-color: #2d3e50 !important;
        color: white !important;
    }
    
    /* Metrics */
    [data-testid="stMetricValue"] {
        color: #c41e3a !important;
    }
</style>
""", unsafe_allow_html=True)

st.title("📝 Report Evaluation")
st.markdown(f"**Welcome, {st.session_state.user_name}!**")

# Sidebar
with st.sidebar:
    st.header("🎯 Quick Actions")
    st.markdown("### 📚 Features")
    if st.button("📸 Bluebook Marks Extraction", use_container_width=True, key="sidebar_bluebook"):
        st.switch_page("pages/5_📸_Bluebook_Extraction.py")
    st.info("📝 **Report Evaluation**\n\n*You are here*")
    
    st.markdown("---")
    if st.button("🏠 Dashboard", use_container_width=True):
        st.switch_page("pages/4_👨‍🏫_Teacher_Dashboard.py")
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.clear()
        st.switch_page("EduLens.py")

# Main tabs
tab1, tab2, tab3 = st.tabs(["📝 Create Rubric", "📊 View Rubrics & Submissions", "📈 Submission Stats"])

# TAB 1: Create Rubric
with tab1:
    st.header("📝 Create New Rubric Set")
    st.info("ℹ️ Upload a PDF containing your grading rubric. AI will extract the criteria automatically.")
    
    with st.form("rubric_upload_form"):
        uploaded_rubric = st.file_uploader(
            "Upload Rubric PDF",
            type=['pdf'],
            help="Select a PDF file containing your rubric"
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            deadline_date = st.date_input("Submission Deadline (Optional)")
            deadline_time = st.time_input("Deadline Time (IST)")
        
        with col2:
            max_attempts = st.number_input(
                "Max Attempts per Student",
                min_value=1,
                max_value=10,
                value=3,
                help="Leave blank for unlimited"
            )
            unlimited = st.checkbox("Unlimited Attempts")
        
        submitted = st.form_submit_button("🚀 Create Rubric Set", use_container_width=True, type="primary")
        
        if submitted:
            if not uploaded_rubric:
                st.error("⚠️ Please upload a rubric PDF file")
            else:
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                    tmp_file.write(uploaded_rubric.getbuffer())
                    temp_path = tmp_file.name
                
                deadline_iso = None
                deadline_display_str = "None (Unlimited)"
                
                if deadline_date:
                    try:
                        deadline_dt = datetime.combine(deadline_date, deadline_time)
                        deadline_dt = deadline_dt.replace(tzinfo=get_ist_timezone())
                        deadline_iso = deadline_dt.isoformat()
                        deadline_display_str = deadline_dt.strftime('%Y-%m-%d %H:%M:%S %Z')
                    except:
                        pass
                
                final_max_attempts = None if unlimited else max_attempts
                attempts_str = "Unlimited" if unlimited else str(max_attempts)
                
                with st.spinner("🧠 AI is extracting rubric criteria... Please wait..."):
                    try:
                        result = extract_and_save_rubric_from_pdf(
                            temp_path,
                            st.session_state.user_id,
                            deadline_iso,
                            final_max_attempts
                        )
                        
                        if result.get('error'):
                            st.error(f"❌ Error: {result['error']}")
                        else:
                            st.success("✅ Rubric setup complete.")
                            st.balloons()
                            
                            rubric_set_id = result['rubric_set_id']
                            parsed_rubrics = result['parsed_rubrics']
                            
                            # --- LOGIC PARITY: Display exact same confirmation info as main.py ---
                            st.markdown("### 📝 Rubric Setup Summary")
                            st.info(f"""
**Rubric Set ID:** `{rubric_set_id}`  
**Deadline (IST):** {deadline_display_str}  
**Max Attempts:** {attempts_str}  
**Parsed Criteria:** {len(parsed_rubrics)} criteria
""")
                            
                            st.warning("📌 Share this Rubric Set ID with students so they can submit against it.")
                            
                            st.markdown("### 🧠 Extracted Logic (Raw JSON):")
                            st.json(parsed_rubrics, expanded=True)
                        
                        os.unlink(temp_path)
                        
                    except Exception as e:
                        st.error(f"❌ Failed to create rubric: {str(e)}")
                        if os.path.exists(temp_path):
                            os.unlink(temp_path)

# TAB 2: View Rubrics & Submissions
with tab2:
    st.header("📊 All Rubric Sets & Submissions")
    
    rubrics = list_rubric_sets()
    
    if not rubrics:
        st.info("📭 No rubrics created yet. Create your first rubric in the 'Create Rubric' tab!")
    else:
        for rubric in rubrics:
            rubric_id = rubric.get('rubric_set_id', 'N/A')
            
            with st.expander(f"📋 Rubric ID: {rubric_id[:20]}...", expanded=False):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    deadline = rubric.get('deadline', 'None')
                    if deadline and deadline != 'None':
                        try:
                            dt = datetime.fromisoformat(deadline)
                            st.markdown(f"**⏰ Deadline:** {dt.strftime('%Y-%m-%d %H:%M')}")
                        except:
                            st.markdown(f"**⏰ Deadline:** {deadline}")
                    else:
                        st.markdown("**⏰ Deadline:** No deadline")
                
                with col2:
                    max_att = rubric.get('max_attempts')
                    st.markdown(f"**🔁 Max Attempts:** {max_att if max_att else 'Unlimited'}")
                
                with col3:
                    parsed_rubrics = rubric.get('parsed_rubrics', [])
                    st.markdown(f"**📝 Criteria Count:** {len(parsed_rubrics)}")
                
                st.markdown("---")
                st.markdown("### 📋 Rubric Criteria")
                
                if parsed_rubrics:
                    for i, criterion in enumerate(parsed_rubrics, 1):
                        title = criterion.get('title', 'Untitled')
                        desc = criterion.get('description', 'No description')
                        score = criterion.get('max_score', 10)
                        
                        with st.container():
                            st.markdown(f"**{i}. {title}** (Max Score: {score})")
                            st.caption(desc)
                else:
                    st.warning("⚠️ No criteria found for this rubric.")
                
                st.markdown("---")
                st.markdown("### 📊 Student Submissions")
                
                submissions = list_submissions_for_rubric(rubric_id)
                
                if not submissions:
                    st.info("📭 No submission details available.")
                else:
                    sub_data = []
                    for sub in submissions:
                        sub_data.append({
                            "Student ID": sub.get('student_id', 'N/A'),
                            "Attempt": sub.get('attempt_number', 1),
                            "Score": sub.get('result', {}).get('total_score', 0),
                            "Date": sub.get('timestamp', 'N/A')[:16].replace('T', ' ')
                        })
                    
                    st.dataframe(pd.DataFrame(sub_data), use_container_width=True)

# TAB 3: Submission Stats
with tab3:
    st.header("📈 Submission Statistics & Reports")
    
    rubrics = list_rubric_sets()
    
    if not rubrics:
        st.info("📭 No rubrics to analyze.")
    else:
        # Create a mapping for dropdown
        rubric_options = {r['rubric_set_id']: f"{r['rubric_set_id'][:15]}... ({len(r.get('submissions_summary', []))} submissions)" for r in rubrics}
        
        selected_rubric_id = st.selectbox(
            "Select Rubric Set",
            options=list(rubric_options.keys()),
            format_func=lambda x: rubric_options[x]
        )
        
        if selected_rubric_id:
            # Find selected rubric object
            selected_rubric = next((r for r in rubrics if r['rubric_set_id'] == selected_rubric_id), None)
            
            if selected_rubric:
                summary_data = selected_rubric.get('submissions_summary', [])
                
                if summary_data:
                    # Metrics
                    total_subs = len(summary_data)
                    avg_score = sum(s.get('total_score', 0) for s in summary_data) / total_subs if total_subs else 0
                    
                    m1, m2 = st.columns(2)
                    m1.metric("Total Submissions", total_subs)
                    m2.metric("Average Score", f"{avg_score:.2f}")
                    
                    # Display Table
                    st.markdown("### 📋 Submission Records")
                    df = pd.DataFrame(summary_data)
                    
                    # Reorder/Rename columns for display
                    display_cols = ['student_id', 'project_name', 'total_score', 'timestamp']
                    df_display = df[display_cols].copy()
                    df_display.columns = ['Student ID', 'Project Name', 'Total Score', 'Timestamp']
                    
                    st.dataframe(df_display, use_container_width=True)
                    
                    # Downloads
                    st.markdown("### 💾 Export Report")
                    
                    col_d1, col_d2 = st.columns(2)
                    
                    # Generate PDF Data
                    try:
                        from fpdf import FPDF
                        import tempfile
                        
                        class PDF(FPDF):
                            def header(self):
                                self.set_font('Helvetica', 'B', 15)
                                self.cell(0, 10, 'Submission Report', align='C')
                                self.ln(20)
                                
                            def footer(self):
                                self.set_y(-15)
                                self.set_font('Helvetica', 'I', 8)
                                self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

                        pdf = PDF()
                        pdf.alias_nb_pages()
                        pdf.add_page()
                        pdf.set_font("Helvetica", size=12)
                        
                        # Rubric Info
                        pdf.set_font("Helvetica", 'B', 12)
                        pdf.cell(0, 10, f"Rubric ID: {selected_rubric_id}", ln=True)
                        pdf.cell(0, 10, f"Total Submissions: {total_subs}", ln=True)
                        pdf.cell(0, 10, f"Average Score: {avg_score:.2f}", ln=True)
                        pdf.ln(10)
                        
                        # Table Header
                        pdf.set_font("Helvetica", 'B', 10)
                        pdf.set_fill_color(200, 220, 255)
                        pdf.cell(40, 10, "Student ID", border=1, fill=True)
                        pdf.cell(90, 10, "Project Name", border=1, fill=True)
                        pdf.cell(30, 10, "Score", border=1, fill=True)
                        pdf.ln()
                        
                        # Table Rows
                        pdf.set_font("Helvetica", size=10)
                        for item in summary_data:
                            student = str(item.get('student_id', 'N/A'))
                            proj = str(item.get('project_name', 'N/A'))[:50]
                            score = str(item.get('total_score', 0))
                            
                            pdf.cell(40, 10, student, border=1)
                            pdf.cell(90, 10, proj, border=1)
                            pdf.cell(30, 10, score, border=1)
                            pdf.ln()
                        
                        # Output PDF as bytes
                        # FPDF's output(dest='S') returns a string, we need bytes.
                        # Using tempfile is safer for FPDF compatibility.
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
                            pdf.output(tmp_pdf.name)
                            tmp_pdf_path = tmp_pdf.name
                            
                        with open(tmp_pdf_path, "rb") as f:
                            pdf_bytes = f.read()
                        os.unlink(tmp_pdf_path)
                        
                        with col_d1:
                            st.download_button(
                                label="📥 Download PDF",
                                data=pdf_bytes,
                                file_name=f"report_{selected_rubric_id[:8]}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                            
                    except Exception as e:
                        with col_d1:
                            st.error(f"PDF Error: {str(e)}")

                    # Generate DOCX Data
                    try:
                        from docx import Document
                        from docx.shared import Inches
                        import io
                        
                        doc = Document()
                        doc.add_heading('Submission Report', 0)
                        
                        doc.add_paragraph(f"Rubric ID: {selected_rubric_id}")
                        doc.add_paragraph(f"Total Submissions: {total_subs}")
                        doc.add_paragraph(f"Average Score: {avg_score:.2f}")
                        
                        table = doc.add_table(rows=1, cols=3)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Student ID'
                        hdr_cells[1].text = 'Project Name'
                        hdr_cells[2].text = 'Score'
                        
                        for item in summary_data:
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(item.get('student_id', 'N/A'))
                            row_cells[1].text = str(item.get('project_name', 'N/A'))
                            row_cells[2].text = str(item.get('total_score', 0))
                            
                        # Save to bytes
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        with col_d2:
                            st.download_button(
                                label="📥 Download Document",
                                data=doc_io,
                                file_name=f"report_{selected_rubric_id[:8]}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                            
                    except Exception as e:
                        with col_d2:
                             st.error(f"DOCX Error: {str(e)}")
                            
                else:
                    st.info("ℹ️ No submissions recorded yet for this rubric.")

